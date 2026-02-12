import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from io import BytesIO
import calendar
from datetime import datetime, timedelta, timezone
import os
import base64

# --- BIBLIOTECAS DE E-MAIL (PADRÃO PYTHON) ---
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# ==============================================================================
# ⚙️ CONFIGURAÇÃO DE E-MAIL (CONFIGURE AQUI)
# ==============================================================================
# 1. E-mail que vai ENVIAR (Seu Gmail ou da Escola)
EMAIL_REMETENTE = "seu_email_aqui@gmail.com" 

# 2. Senha de App (Gerada na segurança do Google - 16 letras)
SENHA_APP_GOOGLE = "xxxx xxxx xxxx xxxx" 

# 3. E-mail da Coordenação que receberá sempre
EMAIL_COORDENACAO = "coordenacao.ceief@gmail.com" 
# ==============================================================================

# --- MATRIZ CURRICULAR ---
try:
    from dados_curriculo import CURRICULO_DB
except ModuleNotFoundError:
    st.error("ERRO: Base de dados curricular não encontrada.")
    st.stop()

# --- 1. CONFIGURAÇÃO DE ALTA PERFORMANCE ---
st.set_page_config(
    page_title="Sistema Planejar | CEIEF",
    layout="wide",
    page_icon="🎓",
    initial_sidebar_state="collapsed"
)

# --- 2. GESTÃO DE ESTADO ---
if 'step' not in st.session_state: st.session_state.step = 1
if 'conteudos_selecionados' not in st.session_state: st.session_state.conteudos_selecionados = []
if 'config' not in st.session_state: st.session_state.config = {}

def set_step(s): st.session_state.step = s

# --- 3. ESTILIZAÇÃO CSS (PREMIUM UI) ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Outfit', sans-serif;
        color: #1e293b;
    }
    
    .stApp { background-color: #f8fafc; }
    [data-testid="stSidebar"], [data-testid="stSidebarNav"] { display: none !important; }
    .st-emotion-cache-16ids0d { display: none !important; }
    
    .block-container { padding-top: 1rem !important; max-width: 1100px !important; }

    /* QUADRANTE DO LOGO */
    .logo-quadrant {
        display: flex; align-items: center; justify-content: center;
        background: white; padding: 10px; border-radius: 18px;
        border: 2px solid #e2e8f0; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
        height: 130px;
    }

    /* HEADER */
    .premium-header-box {
        background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%);
        padding: 2.2rem; border-radius: 20px;
        box-shadow: 0 10px 25px -5px rgba(30, 58, 138, 0.3);
        text-align: center; border: 1px solid rgba(255,255,255,0.1);
        height: 130px;
        display: flex; flex-direction: column; justify-content: center;
    }
    .header-text-main { margin: 0; font-weight: 800; font-size: 2.5rem !important; color: white !important; letter-spacing: -1.5px; line-height: 1; }
    .header-text-sub { margin: 8px 0 0 0; font-weight: 400; color: rgba(255,255,255,0.9); font-size: 1rem; text-transform: uppercase; letter-spacing: 1px; }

    /* CARDS E INPUTS */
    .card-container { background: white; border-radius: 16px; padding: 2.5rem; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); border: 1px solid #e2e8f0; margin-bottom: 1.5rem; }
    .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] {
        border: 2px solid #cbd5e1 !important; border-radius: 12px !important;
        background-color: #ffffff !important; color: #0f172a !important; font-weight: 500 !important;
    }
    
    /* BOTÕES */
    .stButton > button { border-radius: 12px; height: 3.8rem; font-weight: 700; font-size: 1.1rem; text-transform: uppercase; letter-spacing: 0.5px; border: none; transition: all 0.2s ease; }
    div[data-testid="stVerticalBlock"] > div > div > div > div > button[kind="primary"] { background: #1e3a8a !important; color: white !important; }

    /* TAGS */
    .status-tag { display: inline-block; padding: 6px 16px; border-radius: 8px; font-size: 0.75rem; font-weight: 700; text-transform: uppercase; margin-bottom: 12px; border: 1px solid transparent; }
    .tag-blue { background-color: #eff6ff; color: #1e40af; border-color: #bfdbfe; }
    .tag-green { background-color: #f0fdf4; color: #166534; border-color: #bbf7d0; }
    .tag-orange { background-color: #fffaf2; color: #9a3412; border-color: #fed7aa; }

    @media (max-width: 768px) {
        .premium-header-box { height: auto; padding: 1.5rem; }
        .header-text-main { font-size: 1.8rem !important; }
        .logo-quadrant { height: 80px; margin-top: 10px; }
        .logo-pencil { font-size: 2rem !important; }
    }
</style>
""", unsafe_allow_html=True)

# --- FUNÇÕES DE APOIO ---
def get_brazil_time():
    fuso_br = timezone(timedelta(hours=-3))
    return datetime.now(fuso_br)

def clean(t): 
    return t.encode('latin-1', 'replace').decode('latin-1') if t else ""

# --- FUNÇÃO DE E-MAIL AUTOMÁTICO ---
def enviar_email_automatico(pdf_bytes, dados, nome_arquivo):
    """Envia o PDF para a coordenação e CC para o professor."""
    # Validação básica de configuração
    if "xxxx" in SENHA_APP_GOOGLE:
        return False, "⚠️ Configuração de e-mail pendente (Senha de App)."
    
    try:
        msg = MIMEMultipart()
        msg['From'] = EMAIL_REMETENTE
        msg['To'] = EMAIL_COORDENACAO
        
        # Lógica de Cópia (Cc) para o Professor
        destinatarios = [EMAIL_COORDENACAO]
        if dados.get('email_prof') and "@" in dados['email_prof']:
            msg['Cc'] = dados['email_prof']
            destinatarios.append(dados['email_prof'])

        msg['Subject'] = f"Planejamento Entregue: {dados['professor']} - {dados['mes']}"

        corpo = f"""
        Olá,

        Um novo planejamento foi gerado e entregue pelo Sistema Planejar Elite.

        DADOS DO REGISTRO:
        -----------------------------------
        Professor(a): {dados['professor']}
        Ano/Turma: {dados['ano']} - {', '.join(dados['turmas'])}
        Período: {dados['periodo']} ({dados['quinzena']})
        Data de Emissão: {get_brazil_time().strftime("%d/%m/%Y às %H:%M")}
        -----------------------------------

        O documento PDF segue em anexo para validação da coordenação e arquivo do professor.
        
        Atenciosamente,
        Sistema Planejar Elite
        """
        msg.attach(MIMEText(corpo, 'plain'))

        # Anexo
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(pdf_bytes)
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f"attachment; filename= {nome_arquivo}.pdf")
        msg.attach(part)

        # Envio SMTP Gmail
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(EMAIL_REMETENTE, SENHA_APP_GOOGLE)
        server.sendmail(EMAIL_REMETENTE, destinatarios, msg.as_string())
        server.quit()
        return True, "E-mail enviado com sucesso para Coordenação e Professor!"
    except Exception as e:
        return False, f"Falha no envio do e-mail: {str(e)}"

# --- 4. RENDERIZAÇÃO DO CABEÇALHO ---
st.write("") 
col_main, col_logo = st.columns([8, 2], vertical_alignment="center")

with col_main:
    st.markdown(f"""
    <div class="premium-header-box">
        <h1 class="header-text-main">Sistema Planejar</h1>
        <p class="header-text-sub">Gestão Pedagógica Digital • CEIEF Rafael Affonso Leite</p>
    </div>
    """, unsafe_allow_html=True)

with col_logo:
    st.markdown("""<div class="logo-quadrant"><div class="logo-pencil" style="font-size: 3.8rem; text-align: center;">✏️</div></div>""", unsafe_allow_html=True)

# --- NAVEGAÇÃO ---
st.write("")
progresso = {1: 33, 2: 66, 3: 100}
st.progress(progresso[st.session_state.step])
st.write("")

# --- PASSO 1: IDENTIFICAÇÃO ---
if st.session_state.step == 1:
    st.markdown('<div class="card-container">', unsafe_allow_html=True)
    st.markdown("### 📋 Identificação do Planejamento")
    st.write("")
    
    # Linha 1: Professor e Email
    c1, c2 = st.columns(2)
    with c1:
        professor = st.text_input("PROFESSOR(A) RESPONSÁVEL", value=st.session_state.config.get('professor', ''), placeholder="Nome Completo")
    with c2:
        email_prof = st.text_input("E-MAIL DO PROFESSOR (Para receber cópia)", value=st.session_state.config.get('email_prof', ''), placeholder="exemplo@email.com")

    st.write("") # Espaço

    # Linha 2: Ano e Turmas
    c3, c4 = st.columns(2)
    with c3:
        anos = list(CURRICULO_DB.keys())
        if "Maternal I" in anos: anos.remove("Maternal I"); anos.insert(0, "Maternal I")
        saved_ano = st.session_state.config.get('ano')
        idx_ano = anos.index(saved_ano) if saved_ano in anos else 0
        ano = st.selectbox("ANO DE ESCOLARIDADE", anos, index=idx_ano)
        
        if "Maternal" in ano: opts = [f"{ano} - Turma 1", f"{ano} - Turma 2"]
        else:
            qtd = {"Etapa I": 3, "Etapa II": 3, "1º Ano": 3, "2º Ano": 3, "3º Ano": 3, "4º Ano": 3, "5º Ano": 3}
            max_t = qtd.get(ano, 3)
            opts = [f"{prefix}{i}" for i in range(1, max_t + 1) for prefix in ([f"{ano} - Turma " if "Etapa" in ano else f"{ano} "])]
        turmas = st.multiselect("TURMAS VINCULADAS", opts, default=[t for t in st.session_state.config.get('turmas', []) if t in opts])
    
    with c4:
        meses = {2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"}
        saved_mes = st.session_state.config.get('mes')
        idx_mes = list(meses.values()).index(saved_mes) if saved_mes in list(meses.values()) else 0
        mes_nome = st.selectbox("MÊS DE REFERÊNCIA", list(meses.values()), index=idx_mes)
        mes_num = [k for k, v in meses.items() if v == mes_nome][0]
        
        if mes_num == 2:
            quinzena_label, periodo_texto, trimestre_doc = "Mês Inteiro", "01/02/2026 a 28/02/2026", "1º Trimestre"
            st.info("Nota: Fevereiro é Planejamento Mensal.")
        else:
            q_sel = st.radio("PERÍODO DE EXECUÇÃO", ["1ª Quinzena (01-15)", "2ª Quinzena (16-Fim)"], horizontal=True)
            quinzena_label = q_sel.split(" (")[0]
            tri = "1º Trimestre" if mes_num <= 4 else "2º Trimestre" if mes_num <= 8 else "3º Trimestre"
            ultimo = calendar.monthrange(2026, mes_num)[1]
            periodo_texto = f"01/{mes_num:02d}/2026 a 15/{mes_num:02d}/2026" if "1ª" in q_sel else f"16/{mes_num:02d}/2026 a {ultimo}/{mes_num:02d}/2026"
            trimestre_doc = tri

    st.markdown('</div>', unsafe_allow_html=True)
    
    if st.button("Avançar para Matriz Curricular ➔", type="primary", use_container_width=True):
        if not professor or not turmas or not email_prof:
            st.error("ERRO: Todos os campos (incluindo e-mail) são obrigatórios.")
        else:
            if st.session_state.config.get('ano') != ano: st.session_state.conteudos_selecionados = []
            st.session_state.config = {
                'professor': professor, 'email_prof': email_prof, 'ano': ano, 'turmas': turmas, 
                'mes': mes_nome, 'periodo': periodo_texto, 'trimestre': trimestre_doc, 'quinzena': quinzen_label if 'quinzena_label' in locals() else quinzen_label
            }
            set_step(2); st.rerun()

# --- PASSO 2: MATRIZ ---
elif st.session_state.step == 2:
    ano_sel = st.session_state.config['ano']
    st.markdown(f"### 📖 Matriz Curricular: **{ano_sel}**")
    with st.container():
        st.markdown('<div class="card-container">', unsafe_allow_html=True)
        dados = CURRICULO_DB.get(ano_sel, {})
        infantil_anos = ["Maternal I", "Maternal II", "Etapa I", "Etapa II"]
        
        if ano_sel in infantil_anos:
            abas = st.tabs(["🗣️ Linguagem Verbal", "🔢 Linguagem Matemática", "👥 Indivíduo e Sociedade"])
            chaves = ["LINGUAGEM VERBAL", "LINGUAGEM MATEMÁTICA", "INDIVÍDUO E SOCIEDADE"]
            tags = ["tag-blue", "tag-orange", "tag-green"]
        else:
            abas = st.tabs(["💻 Tecnologia & Cultura Digital", "🗣️ Língua Inglesa"])
            op_tec = [k for k, v in dados.items() if "INGLÊS" not in k.upper() and (v and "ORALIDADE" not in v[0]['eixo'].upper())]
            op_ing = [k for k in dados.keys() if k not in op_tec]
            chaves = [op_tec, op_ing]
            tags = ["tag-blue", "tag-green"]

        for idx, aba in enumerate(abas):
            with aba:
                if ano_sel in infantil_anos:
                    area = chaves[idx]
                    if area in dados:
                        c1, c2 = st.columns(2)
                        opcoes_g = sorted(list(set([it['geral'] for it in dados[area]])))
                        g_sel = c1.selectbox(f"CONTEÚDO GERAL", opcoes_g, key=f"inf_g_{idx}")
                        filtro = [it for it in dados[area] if it['geral'] == g_sel]
                        es = [it['especifico'] for it in filtro]
                        e_sel = c2.selectbox(f"CONTEÚDO ESPECÍFICO", es, key=f"inf_e_{idx}")
                        sel = next((it for it in filtro if it['especifico'] == e_sel), None)
                        if sel:
                            st.markdown(f"<div style='background:#f8fafc; padding:1.2rem; border-radius:12px; border:1px solid #cbd5e1; margin-top:10px;'><span class='status-tag {tags[idx]}'>Objetivo Pedagógico</span><br><b>{sel['objetivo']}</b></div>", unsafe_allow_html=True)
                            if st.button("Adicionar à Lista ➕", key=f"btn_inf_{idx}"):
                                st.session_state.conteudos_selecionados.append({'tipo': area, 'eixo': sel['eixo'], 'geral': g_sel, 'especifico': e_sel, 'objetivo': sel['objetivo']})
                                st.toast("Item adicionado!")
                else:
                    filtros = chaves[idx]
                    if filtros:
                        c1, c2 = st.columns(2)
                        g = c1.selectbox("EIXO / TÓPICO", filtros, key=f"f_g_{idx}")
                        e = c2.selectbox("CONTEÚDO / PRÁTICA", [it['especifico'] for it in dados[g]], key=f"f_e_{idx}")
                        sel = next((it for it in dados[g] if it['especifico'] == e), None)
                        if sel:
                            st.markdown(f"<div style='background:#f8fafc; padding:1.2rem; border-radius:12px; border:1px solid #cbd5e1; margin-top:10px;'><span class='status-tag {tags[idx]}'>Objetivo do Currículo</span><br><b>{sel['objetivo']}</b></div>", unsafe_allow_html=True)
                            if st.button("Adicionar à Lista ➕", key=f"btn_f_{idx}"):
                                label_tipo = "Tecnologia" if idx == 0 else "Inglês"
                                st.session_state.conteudos_selecionados.append({'tipo': label_tipo, 'eixo': sel['eixo'], 'geral': g, 'especifico': e, 'objetivo': sel['objetivo']})
                                st.toast("Item adicionado!")
        st.markdown('</div>', unsafe_allow_html=True)

    if st.session_state.conteudos_selecionados:
        st.markdown("#### Conteúdos Selecionados")
        for i, it in enumerate(st.session_state.conteudos_selecionados):
            col_t, col_b = st.columns([0.96, 0.04])
            with col_t: st.markdown(f"<div style='background:white; border:1px solid #e2e8f0; padding:1rem; border-radius:12px; margin-bottom:5px;'><b>[{it['tipo']}]</b> {it['geral']}: {it['especifico']}</div>", unsafe_allow_html=True)
            if col_b.button("✕", key=f"del_{i}"): st.session_state.conteudos_selecionados.pop(i); st.rerun()

    c1, c2 = st.columns(2)
    c1.button("⬅ Voltar", on_click=set_step, args=(1,))
    if c2.button("Avançar para Detalhamento ➔", type="primary", use_container_width=True):
        if not st.session_state.conteudos_selecionados: st.error("Selecione um conteúdo.")
        else: set_step(3); st.rerun()

# --- PASSO 3: DETALHAMENTO ---
elif st.session_state.step == 3:
    st.markdown("### ✍️ Detalhamento Pedagógico")
    with st.container():
        st.markdown('<div class="card-container">', unsafe_allow_html=True)
        st.markdown("<div style='color:#be123c; font-weight:800; font-size:0.8rem; margin-bottom:1.5rem;'>PREENCHIMENTO OBRIGATÓRIO</div>", unsafe_allow_html=True)
        
        obj_esp = st.text_area("Objetivos Específicos", height=100, placeholder="Defina os resultados práticos pretendidos...", value=st.session_state.config.get('obj_esp', ''))
        c1, c2 = st.columns(2)
        with c1: sit = st.text_area("Situação didática", height=220, placeholder="Passo a passo...", value=st.session_state.config.get('sit', ''))
        with c2: rec = st.text_area("Recursos e Materiais", height=220, value=st.session_state.config.get('rec', 'Descritos na situação didática'))
        recup = st.text_area("Recuperação Contínua", height=100, value=st.session_state.config.get('recup', ''))
        st.markdown('</div>', unsafe_allow_html=True)

    st.session_state.config.update({'obj_esp': obj_esp, 'sit': sit, 'rec': rec, 'recup': recup})

    def gerar_pdf(dados, conteudos):
        pdf = FPDF(); pdf.add_page(); pdf.set_auto_page_break(auto=True, margin=30)
        logo_e = "logo_escola.png" if os.path.exists("logo_escola.png") else "logo_escola.jpg"
        if os.path.exists(logo_e): pdf.image(logo_e, 175, 8, 25)
        pdf.set_font('Arial', 'B', 14); pdf.cell(0, 10, clean('CEIEF RAFAEL AFFONSO LEITE'), 0, 1, 'C')
        pdf.set_font('Arial', '', 10); pdf.cell(0, 5, clean('Planejamento de Unidade de Ensino'), 0, 1, 'C'); pdf.ln(10)
        pdf.set_fill_color(245, 247, 250); pdf.set_font("Arial", 'B', 9)
        pdf.cell(0, 7, clean(f"DOCENTE: {dados['professor']}"), 1, 1, 'L', True)
        pdf.cell(0, 7, clean(f"ANO: {dados['ano']} | TURMAS: {', '.join(dados['turmas'])}"), 1, 1, 'L', True)
        pdf.cell(0, 7, clean(f"MES: {dados['mes']} | PERIODO: {dados['quinzena']} | TRIMESTRE: {dados['trimestre']}"), 1, 1, 'L', True)
        pdf.cell(0, 7, clean(f"INTERVALO: {dados['periodo']}"), 1, 1, 'L', True); pdf.ln(5)
        pdf.set_font("Arial", 'B', 10); pdf.cell(0, 8, clean("MATRIZ CURRICULAR SELECIONADA"), 0, 1)
        pdf.set_fill_color(230, 230, 230); pdf.set_font("Arial", 'B', 8)
        col_w = [45, 75, 70]
        pdf.cell(col_w[0], 7, clean("Eixo / Tema"), 1, 0, 'C', True); pdf.cell(col_w[1], 7, clean("Habilidade Especifica"), 1, 0, 'C', True); pdf.cell(col_w[2], 7, clean("Objetivo do Ano"), 1, 1, 'C', True)
        pdf.set_font("Arial", '', 8)
        for it in conteudos:
            x, y = pdf.get_x(), pdf.get_y()
            pdf.multi_cell(col_w[0], 5, clean(f"{it['eixo']}\n({it['geral']})"), 0, 'L')
            y1 = pdf.get_y(); pdf.set_xy(x + col_w[0], y)
            pdf.multi_cell(col_w[1], 5, clean(it['especifico']), 0, 'L')
            y2 = pdf.get_y(); pdf.set_xy(x + col_w[0] + col_w[1], y)
            pdf.multi_cell(col_w[2], 5, clean(it['objetivo']), 0, 'L')
            y3 = pdf.get_y(); max_y = max(y1, y2, y3); h_row = max_y - y
            pdf.set_xy(x, y); pdf.cell(col_w[0], h_row, "", 1, 0); pdf.cell(col_w[1], h_row, "", 1, 0); pdf.cell(col_w[2], h_row, "", 1, 1)
            pdf.set_y(max_y)
        pdf.ln(5); pdf.set_font("Arial", 'B', 10); pdf.cell(0, 8, clean("DETALHAMENTO PEDAGOGICO"), 0, 1)
        for l, v in [("Objetivos Especificos", dados['obj_esp']), ("Situação didática", dados['sit']), ("Recursos e Materiais", dados['rec']), ("Recuperação Contínua", dados['recup'])]:
            pdf.set_font("Arial", 'B', 9); pdf.cell(0, 5, clean(l + ":"), 0, 1); pdf.set_font("Arial", '', 9); pdf.multi_cell(0, 5, clean(v)); pdf.ln(2)
        pdf.set_font("Arial", 'B', 9); pdf.cell(0, 5, clean("Avaliação:"), 0, 1)
        pdf.line(pdf.get_x(), pdf.get_y()+5, 200, pdf.get_y()+5); pdf.line(pdf.get_x(), pdf.get_y()+12, 200, pdf.get_y()+12); pdf.ln(15)
        pdf.set_auto_page_break(False); pdf.set_y(-15); pdf.set_font('Arial', 'I', 7)
        pdf.cell(0, 10, clean(f'Emitido via Sistema Planejar em: {get_brazil_time().strftime("%d/%m/%Y %H:%M:%S")} (GMT-3)'), 0, 0, 'C')
        pdf.set_auto_page_break(True, margin=30)
        return bytes(pdf.output(dest='S').encode('latin-1'))

    def gerar_docx(dados, conteudos):
        doc = Document(); style = doc.styles['Normal']; font = style.font; font.name = 'Arial'; font.size = Pt(10)
        table_h = doc.add_table(rows=1, cols=2); table_h.autofit = False; table_h.columns[0].width = Cm(14); table_h.columns[1].width = Cm(4)
        p = table_h.cell(0,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.add_run("CEIEF RAFAEL AFFONSO LEITE\n").bold = True; p.add_run("Planejamento Digital de Linguagens e Tecnologias")
        logo_e = "logo_escola.png" if os.path.exists("logo_escola.png") else "logo_escola.jpg"
        if os.path.exists(logo_e): table_h.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; table_h.cell(0,1).paragraphs[0].add_run().add_picture(logo_e, width=Cm(3.0))
        doc.add_paragraph(); p_info = doc.add_paragraph(); p_info.add_run(f"DOCENTE: {dados['professor']}\n").bold = True; p_info.add_run(f"ANO: {dados['ano']} | TURMAS: {', '.join(dados['turmas'])}\n"); p_info.add_run(f"MES: {dados['mes']} | PERIODO: {dados['quinzena']} | TRIMESTRE: {dados['trimestre']}\n"); p_info.add_run(f"INTERVALO: {dados['periodo']}")
        doc.add_heading("Matriz Curricular Selecionada", 2); table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
        hdr = table.rows[0].cells; hdr[0].text = 'Eixo / Tema'; hdr[1].text = 'Habilidade Especifica'; hdr[2].text = 'Objetivo do Ano'
        for cell in hdr: cell.paragraphs[0].runs[0].bold = True
        for it in conteudos:
            row = table.add_row().cells; row[0].text = f"{it['eixo']}\n({it['geral']})"; row[1].text = it['especifico']; row[2].text = it['objetivo']
        doc.add_heading("Detalhamento Pedagogico", 2)
        for l, v in [("Objetivos Especificos", dados['obj_esp']), ("Situação didática", dados['sit']), ("Recursos e Materiais", dados['rec']), ("Recuperação Contínua", dados['recup'])]:
            p = doc.add_paragraph(); p.add_run(l + ": ").bold = True; p.add_run(v)
        p_aval = doc.add_paragraph(); p_aval.add_run("Avaliação: ").bold = True
        doc.add_paragraph("_" * 80); doc.add_paragraph("_" * 80)
        doc.add_paragraph(f"\nEmitido eletronicamente em: {get_brazil_time().strftime('%d/%m/%Y %H:%M:%S')} (GMT-3)")
        f = BytesIO(); doc.save(f); f.seek(0); return f

    c1, c2 = st.columns(2)
    if c1.button("⬅ Matriz"): set_step(2); st.rerun()
    if c2.button("GERAR PLANEJAMENTO FINAL 🚀", type="primary", use_container_width=True):
        if not all([obj_esp, sit, rec, recup]): st.error("Erro: Preencha todos os campos.")
        else:
            with st.spinner("Gerando documentos e enviando e-mail..."):
                f_data = st.session_state.config
                w_file = gerar_docx(f_data, st.session_state.conteudos_selecionados)
                p_file = gerar_pdf(f_data, st.session_state.conteudos_selecionados)
                nome_arq = f"Plan_{f_data['mes']}_{f_data['ano'].replace(' ','')}"
                
                # TENTA ENVIAR E-MAIL SE O CAMPO ESTIVER PREENCHIDO
                if f_data.get('email_prof'):
                    sucesso_email, msg_email = enviar_email_automatico(p_file, f_data, nome_arq)
                    if sucesso_email: st.success(f"✅ {msg_email}")
                    else: st.warning(f"⚠️ Documentos gerados, mas o e-mail falhou: {msg_email}")
                else:
                    st.info("ℹ️ E-mail não enviado (endereço do professor não informado).")

                cd1, cd2 = st.columns(2)
                cd1.download_button("📄 Download WORD", w_file, f"{nome_arq}.docx", use_container_width=True)
                cd2.download_button("📕 Download PDF", p_file, f"{nome_arq}.pdf", use_container_width=True)

# --- RODAPÉ ---
st.markdown(f"""
    <div style="text-align:center; margin-top:80px; padding:40px; color:#94a3b8; font-size:0.8rem; border-top:1px solid #e2e8f0;">
        <b>SISTEMA PLANEJAR ELITE V9.0</b><br>
        Desenvolvido por José Victor Souza Gallo • CEIEF Rafael Affonso Leite © {datetime.now().year}
    </div>
""", unsafe_allow_html=True)
